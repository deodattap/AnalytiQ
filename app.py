import os, re, math, unicodedata
import numpy as np
import pandas as pd
from collections import Counter
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity
try:
    from mlxtend.frequent_patterns import apriori, association_rules
    import inspect as _insp
    _AR_NUM = 'num_itemsets' in _insp.signature(association_rules).parameters
except ImportError:
    apriori = association_rules = None
    _AR_NUM = False

app = Flask(__name__, template_folder='templates')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
CORS(app)

# ── Skill registry: (canonical, regex, category, weight) ───────────────────────
# weight: 3=core lang, 2=major framework/cloud, 1=tool/soft
_SR = [
    # Programming
    ('python',         r'\bpython\b',                               'pl',  3),
    ('java',           r'\bjava(?!script)\b',                       'pl',  3),
    ('javascript',     r'\bjavascript\b|\bjs\b(?!\s*on)',           'pl',  3),
    ('typescript',     r'\btypescript\b|\bts\b',                    'pl',  3),
    ('c++',            r'\bc\+\+\b|\bcpp\b',                        'pl',  3),
    ('c#',             r'\bc#\b|\bcsharp\b',                        'pl',  3),
    ('c',              r'\bc\b(?!\+\+|#|ss|loud|ard)',              'pl',  2),
    ('golang',         r'\bgolang\b|\bgo\s+(?:lang|programming)\b', 'pl',  3),
    ('rust',           r'\brust\b(?!\s+belt)',                       'pl',  3),
    ('ruby',           r'\bruby\b(?!\s+on\s+rails)',                 'pl',  2),
    ('php',            r'\bphp\b',                                   'pl',  2),
    ('swift',          r'\bswift\b',                                 'pl',  2),
    ('kotlin',         r'\bkotlin\b',                                'pl',  2),
    ('scala',          r'\bscala\b',                                 'pl',  2),
    ('r',              r'\br\b(?=\s*(?:language|programming|studio|shiny))', 'pl', 2),
    ('perl',           r'\bperl\b',                                  'pl',  1),
    ('shell',          r'\bbash\b|\bshell\s*script(?:ing)?\b',       'pl',  2),
    ('matlab',         r'\bmatlab\b',                                'pl',  1),
    ('dart',           r'\bdart\b',                                  'pl',  2),
    ('lua',            r'\blua\b',                                   'pl',  1),
    # Frontend / Web
    ('html',           r'\bhtml5?\b',                                'fe',  2),
    ('css',            r'\bcss3?\b',                                 'fe',  2),
    ('sass',           r'\bsass\b|\bscss\b',                        'fe',  1),
    ('tailwind',       r'\btailwind(?:css)?\b',                     'fe',  1),
    ('bootstrap',      r'\bbootstrap\b',                            'fe',  1),
    ('react',          r'\breact(?:\.js|js)?\b',                    'fe',  3),
    ('angular',        r'\bangular(?:\.?js)?\b',                    'fe',  3),
    ('vue',            r'\bvue(?:\.?js)?\b',                        'fe',  2),
    ('svelte',         r'\bsvelte\b',                               'fe',  2),
    ('next.js',        r'\bnext\.?js\b',                            'fe',  2),
    ('nuxt',           r'\bnuxt(?:\.?js)?\b',                       'fe',  1),
    ('jquery',         r'\bjquery\b',                               'fe',  1),
    ('webpack',        r'\bwebpack\b|\bvite\b',                     'fe',  1),
    ('flutter',        r'\bflutter\b',                              'fe',  2),
    ('react native',   r'\breact\s+native\b',                       'fe',  2),
    # Backend
    ('node.js',        r'\bnode(?:\.?js)?\b',                       'be',  3),
    ('express',        r'\bexpress(?:\.?js)?\b',                    'be',  2),
    ('django',         r'\bdjango\b',                               'be',  3),
    ('flask',          r'\bflask\b',                                'be',  2),
    ('fastapi',        r'\bfastapi\b',                              'be',  2),
    ('spring boot',    r'\bspring\s*boot\b',                        'be',  3),
    ('spring',         r'\bspring\b(?!\s*boot)',                    'be',  2),
    ('laravel',        r'\blaravel\b',                              'be',  2),
    ('rails',          r'\bruby\s+on\s+rails\b|\brails\b',         'be',  2),
    ('asp.net',        r'\basp\.?net\b',                            'be',  2),
    ('graphql',        r'\bgraphql\b',                              'be',  2),
    ('rest api',       r'\brest(?:ful)?\s*api\b|\brest\s+service\b','be',  2),
    ('grpc',           r'\bgrpc\b',                                 'be',  2),
    ('microservices',  r'\bmicroservices?\b',                       'be',  2),
    # Databases
    ('sql',            r'\bsql\b(?!\s*server|\s*noSQL)',            'db',  2),
    ('mysql',          r'\bmysql\b',                                'db',  2),
    ('postgresql',     r'\bpostgresql\b|\bpostgres\b',              'db',  2),
    ('mongodb',        r'\bmongodb\b|\bmongo\b(?!\s*db)',           'db',  2),
    ('redis',          r'\bredis\b',                                'db',  2),
    ('sqlite',         r'\bsqlite\b',                               'db',  1),
    ('firebase',       r'\bfirebase\b',                             'db',  2),
    ('elasticsearch',  r'\belasticsearch\b|\belastic\s+search\b',   'db',  2),
    ('dynamodb',       r'\bdynamodb\b',                             'db',  2),
    ('cassandra',      r'\bcassandra\b',                            'db',  2),
    ('oracle db',      r'\boracle\s*(?:db|database)?\b',            'db',  2),
    ('nosql',          r'\bnosql\b',                                'db',  1),
    ('supabase',       r'\bsupabase\b',                             'db',  1),
    # Cloud & DevOps
    ('aws',            r'\baws\b|\bamazon\s+web\s+services\b',      'cloud', 3),
    ('azure',          r'\b(?:microsoft\s+)?azure\b',               'cloud', 3),
    ('gcp',            r'\bgcp\b|\bgoogle\s+cloud\b',               'cloud', 3),
    ('docker',         r'\bdocker\b',                               'devops', 3),
    ('kubernetes',     r'\bkubernetes\b|\bk8s\b',                   'devops', 3),
    ('terraform',      r'\bterraform\b',                            'devops', 2),
    ('ansible',        r'\bansible\b',                              'devops', 2),
    ('jenkins',        r'\bjenkins\b',                              'devops', 2),
    ('github actions', r'\bgithub\s+actions\b',                     'devops', 2),
    ('ci/cd',          r'\bci\s*/\s*cd\b|\bcontinuous\s+(?:integration|deployment|delivery)\b', 'devops', 2),
    ('linux',          r'\blinux\b|\bubuntu\b|\bcentos\b|\bdebian\b|\barch\s+linux\b', 'devops', 2),
    ('git',            r'\bgit\b(?!\s*hub|\s*lab)',                 'devops', 2),
    ('github',         r'\bgithub\b(?!\s+actions)',                 'devops', 1),
    ('gitlab',         r'\bgitlab\b',                               'devops', 1),
    ('helm',           r'\bhelm\b',                                 'devops', 1),
    ('nginx',          r'\bnginx\b',                                'devops', 1),
    # Data Science / ML / AI
    ('machine learning', r'\bmachine\s+learning\b|\b(?<!\w)ml\b',  'ml', 3),
    ('deep learning',  r'\bdeep\s+learning\b|\b(?<!\w)dl\b',       'ml', 3),
    ('tensorflow',     r'\btensorflow\b|\b(?<!\w)tf\b(?!\s*idf)',   'ml', 3),
    ('pytorch',        r'\bpytorch\b|\btorch\b',                    'ml', 3),
    ('scikit-learn',   r'\bscikit[-\s]?learn\b|\bsklearn\b',        'ml', 2),
    ('pandas',         r'\bpandas\b',                               'ml', 2),
    ('numpy',          r'\bnumpy\b',                                'ml', 2),
    ('data analysis',  r'\bdata\s+anal(?:ysis|ytics)\b',            'ml', 2),
    ('nlp',            r'\bnlp\b|\bnatural\s+language\s+processing\b', 'ml', 3),
    ('computer vision',r'\bcomputer\s+vision\b|\b(?<!\w)cv\b(?=\s+(?:model|engineer|project))', 'ml', 3),
    ('llm',            r'\bllm\b|\blarge\s+language\s+model\b|\bgpt\b|\bopenai\b|\blangchain\b', 'ml', 3),
    ('tableau',        r'\btableau\b',                              'ml', 2),
    ('power bi',       r'\bpower\s*bi\b',                           'ml', 2),
    ('excel',          r'\b(?:microsoft\s+)?excel\b',               'ml', 1),
    ('hadoop',         r'\bhadoop\b|\bhdfs\b',                      'ml', 2),
    ('spark',          r'\b(?:apache\s+)?spark\b|\bpyspark\b',      'ml', 2),
    ('airflow',        r'\b(?:apache\s+)?airflow\b',                'ml', 2),
    ('dbt',            r'\bdbt\b',                                  'ml', 1),
    # Design / UX
    ('figma',          r'\bfigma\b',                                'design', 2),
    ('sketch',         r'\bsketch\b',                               'design', 1),
    ('adobe xd',       r'\badobe\s+xd\b',                          'design', 1),
    ('photoshop',      r'\bphotoshop\b',                            'design', 1),
    ('ux design',      r'\bux\s*(?:design|designer|research)\b|\buser\s+experience\b', 'design', 2),
    ('ui design',      r'\bui\s*(?:design|designer)\b|\buser\s+interface\s+design\b', 'design', 2),
    ('prototyping',    r'\bprototyping?\b',                         'design', 1),
    ('wireframing',    r'\bwireframing?\b',                         'design', 1),
    ('design systems', r'\bdesign\s+systems?\b',                    'design', 1),
    ('a/b testing',    r'\ba/?b\s+test(?:ing)?\b',                  'design', 1),
    # Methodology / Soft
    ('agile',          r'\bagile\b',                                'soft', 1),
    ('scrum',          r'\bscrum\b',                                'soft', 1),
    ('jira',           r'\bjira\b',                                 'soft', 1),
    ('system design',  r'\bsystem\s+design\b',                      'soft', 2),
    ('problem solving',r'\bproblem[\s\-]?solving\b',                'soft', 1),
    ('communication',  r'\bcommunication\b',                        'soft', 1),
    ('leadership',     r'\bleadership\b',                           'soft', 1),
    ('teamwork',       r'\bteamwork\b|\bteam\s+player\b',           'soft', 1),
]

# Pre-compile patterns once at startup
SKILL_PATTERNS = [(name, re.compile(pat, re.I), cat, w) for name, pat, cat, w in _SR]
SKILLS = [name for name, *_ in _SR]  # flat list for vector indexing

SECTION_RE = {
    'skills':   re.compile(r'(?:^|\n)\s*(?:technical\s+)?skills?(?:\s+&\s+\w+)?\s*[:—\n]', re.I),
    'experience': re.compile(r'(?:^|\n)\s*(?:work\s+)?experience|employment|internship\b', re.I),
    'education': re.compile(r'(?:^|\n)\s*education(?:al)?\s*(?:background|qualification)?\s*[:—\n]?', re.I),
    'projects':  re.compile(r'(?:^|\n)\s*projects?\s*[:—\n]', re.I),
    'certifications': re.compile(r'(?:^|\n)\s*certifi(?:cations?|cates?)|awards?\s*[:—\n]', re.I),
    'summary':   re.compile(r'(?:^|\n)\s*(?:professional\s+)?(?:summary|objective|profile|about)\s*[:—\n]?', re.I),
}

EDU_LEVELS = [
    (re.compile(r'\bph\.?d\.?|doctor(?:ate|al)\b', re.I), 5),
    (re.compile(r'\bm\.?\s*(?:tech|sc|s|eng|ba|ca|b\.?e)\b|master', re.I), 4),
    (re.compile(r'\bmba\b', re.I), 4),
    (re.compile(r'\bb\.?\s*(?:tech|sc|s|e|eng|ca)\b|bachelor|undergraduate|b\.?e\b', re.I), 3),
    (re.compile(r'\bdiploma\b|\bpolytechnic\b', re.I), 2),
    (re.compile(r'\bschool\b|\bssc\b|\bhsc\b|\b12th\b|\b10th\b', re.I), 1),
]

CERT_RE = re.compile(
    r'\b(?:aws|azure|gcp|google|oracle|cisco|microsoft|pmp|comptia|certified|'
    r'certification|certificate|cka|ckad|ccp|csa|clf|saa|dva|soa|dop|mls|'
    r'coursera|udemy|edx|leetcode|hackerrank|kaggle)\b',
    re.I
)

SKILL_META = {
    'python':('🐍','linear-gradient(135deg,#3776ab,#ffd343)'),
    'java':('☕','linear-gradient(135deg,#f89820,#5382a1)'),
    'javascript':('⚡','linear-gradient(135deg,#f7df1e,#c8b400)'),
    'typescript':('📘','linear-gradient(135deg,#3178c6,#1e4f8c)'),
    'react':('⚛️','linear-gradient(135deg,#61dafb,#21232a)'),
    'angular':('🅰️','linear-gradient(135deg,#dd0031,#c3002f)'),
    'vue':('🟩','linear-gradient(135deg,#4fc08d,#35495e)'),
    'node.js':('🟩','linear-gradient(135deg,#68a063,#215732)'),
    'django':('🌿','linear-gradient(135deg,#092e20,#44b78b)'),
    'flask':('🌶️','linear-gradient(135deg,#000,#333)'),
    'spring boot':('🍃','linear-gradient(135deg,#6db33f,#3d6b22)'),
    'docker':('🐳','linear-gradient(135deg,#0db7ed,#384d54)'),
    'kubernetes':('☸️','linear-gradient(135deg,#326ce5,#1a3a8f)'),
    'aws':('☁️','linear-gradient(135deg,#ff9900,#c45000)'),
    'azure':('☁️','linear-gradient(135deg,#0078d4,#004578)'),
    'gcp':('☁️','linear-gradient(135deg,#4285f4,#0f9d58)'),
    'machine learning':('🤖','linear-gradient(135deg,#8e44ad,#c0392b)'),
    'deep learning':('🧬','linear-gradient(135deg,#6a11cb,#2575fc)'),
    'tensorflow':('🧠','linear-gradient(135deg,#ff6f00,#ffa800)'),
    'pytorch':('🔥','linear-gradient(135deg,#ee4c2c,#8b1a09)'),
    'sql':('🗄️','linear-gradient(135deg,#f7971e,#ffd200)'),
    'mongodb':('🍃','linear-gradient(135deg,#47a248,#1a5c1e)'),
    'postgresql':('🐘','linear-gradient(135deg,#336791,#1a3a5c)'),
    'redis':('🔴','linear-gradient(135deg,#d82c20,#9c1c14)'),
    'git':('🔀','linear-gradient(135deg,#f05032,#333)'),
    'linux':('🐧','linear-gradient(135deg,#fcc624,#2c2c2c)'),
    'figma':('🎨','linear-gradient(135deg,#f24e1e,#a259ff)'),
    'rest api':('🔗','linear-gradient(135deg,#11998e,#38ef7d)'),
    'graphql':('💡','linear-gradient(135deg,#e10098,#7b0055)'),
    'ci/cd':('⚙️','linear-gradient(135deg,#2c3e50,#4ca1af)'),
    'system design':('🏗️','linear-gradient(135deg,#8e44ad,#c0392b)'),
    'pandas':('🐼','linear-gradient(135deg,#150458,#090226)'),
    'numpy':('🔢','linear-gradient(135deg,#4d77cf,#013220)'),
    'spark':('✨','linear-gradient(135deg,#e25a1c,#f7931e)'),
    'tableau':('📊','linear-gradient(135deg,#1f77b4,#005073)'),
    'power bi':('📊','linear-gradient(135deg,#f2c811,#e07b00)'),
    'nlp':('💬','linear-gradient(135deg,#667eea,#764ba2)'),
    'llm':('🧠','linear-gradient(135deg,#43cea2,#185a9d)'),
    'flutter':('🦋','linear-gradient(135deg,#54c5f8,#01579b)'),
    'react native':('📱','linear-gradient(135deg,#61dafb,#0366d6)'),
}
_DM = {'emoji':'🛠️','bg':'linear-gradient(135deg,#005f98,#2aa7ff)'}

SECTION_KEYWORDS = {
    'skills':['skill','technical skill','core competenc','tech stack','tools','languages','frameworks','expertise','technologies'],
    'experience':['experience','work experience','professional experience','employment','work history','internship','industrial training','career'],
    'education':['education','academic','qualification','degree','bachelor','master','phd','diploma','university','college','b.tech','m.tech','mba'],
    'projects':['project','portfolio','built','developed','created','implemented','github','application'],
    'certifications':['certification','certified','certificate','award','achievement','credential','honour'],
    'summary':['summary','objective','career objective','profile','about me','professional summary','overview'],
}

# ── Text utilities ──────────────────────────────────────────────────────────────

def clean_text(raw: str) -> str:
    text = unicodedata.normalize('NFKD', raw)
    text = text.replace('\x00',' ').replace('\r\n','\n').replace('\r','\n')
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def extract_text(file) -> str:
    fname = file.filename.lower()
    raw = ''
    if fname.endswith('.pdf'):
        try:
            import pdfplumber
            with pdfplumber.open(file) as pdf:
                pages = []
                for pg in pdf.pages:
                    t = pg.extract_text(x_tolerance=3, y_tolerance=3)
                    if t: pages.append(t)
                raw = '\n'.join(pages)
        except Exception as e:
            app.logger.warning(f'[PDF] {file.filename}: {e}')
    elif fname.endswith(('.docx','.doc')):
        try:
            from docx import Document
            doc = Document(file)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text)
            raw = '\n'.join(parts)
        except Exception as e:
            app.logger.warning(f'[DOCX] {file.filename}: {e}')
    return clean_text(raw)

# ── Section splitter ────────────────────────────────────────────────────────────
def split_sections(text: str) -> dict:
    """Split resume text into labeled sections."""
    sections = {k: '' for k in SECTION_KEYWORDS}
    lower = text.lower()
    # Find all section boundaries
    boundaries = []
    for sec, kws in SECTION_KEYWORDS.items():
        for kw in kws:
            for m in re.finditer(r'(?:^|\n)[ \t]*' + re.escape(kw) + r'[^\n]{0,25}\n', lower):
                boundaries.append((m.start(), sec))
    boundaries.sort()
    # Assign text between boundaries
    for i, (pos, sec) in enumerate(boundaries):
        end = boundaries[i+1][0] if i+1 < len(boundaries) else len(text)
        chunk = text[pos:end]
        sections[sec] = sections[sec] + '\n' + chunk if sections[sec] else chunk
    sections['_full'] = text
    return sections

# ── Skill extraction ────────────────────────────────────────────────────────────
def detect_skills(text: str, weighted: bool = False):
    """
    Extract skills from text using pre-compiled regex patterns.
    If weighted=True, returns dict {skill: weight_score} for scoring.
    Otherwise returns ordered list of canonical skill names.
    """
    lower = text.lower()
    found = {}
    seen = set()
    for name, pat, cat, w in SKILL_PATTERNS:
        if name in seen: continue
        if pat.search(lower):
            found[name] = w
            seen.add(name)
    if weighted:
        return found
    return list(found.keys())


def detect_skills_sectioned(text: str) -> dict:
    """
    Section-aware skill extraction with frequency boosting.
    Skills in 'skills' section get highest confidence.
    Skills found ONLY in _full fallback (not any named section) are penalised.
    Returns {skill_name: confidence_score 0-100}.
    """
    secs = split_sections(text)
    skill_confidence = {}
    # Section priority weights (section_name -> multiplier)
    SEC_WEIGHTS = {
        'skills':         1.00,   # explicitly listed skills
        'experience':     0.82,   # used at work
        'projects':       0.88,   # demonstrated in projects
        'certifications': 0.75,   # certified in
        'summary':        0.65,   # mentioned in summary
        'education':      0.55,   # academic
    }
    # Track which sections each skill appears in
    skill_sections: dict[str, list[float]] = {}

    for sec_name, multiplier in SEC_WEIGHTS.items():
        chunk = secs.get(sec_name, '')
        if not chunk:
            continue
        chunk_lower = chunk.lower()
        for name, pat, cat, base_w in SKILL_PATTERNS:
            if pat.search(chunk_lower):
                pts = multiplier * base_w * 33.3   # max ~100 for w=3 in skills
                skill_sections.setdefault(name, []).append(pts)

    # _full fallback: pick up skills missed by section splitter
    full_lower = text.lower()
    for name, pat, cat, base_w in SKILL_PATTERNS:
        if name not in skill_sections and pat.search(full_lower):
            # Found only in full text — low confidence, keep it but penalised
            skill_sections[name] = [base_w * 16.0]   # ~16-48 range

    # Final confidence = best section score + frequency bonus
    for name, scores in skill_sections.items():
        best = max(scores)
        freq_bonus = min(10, (len(scores) - 1) * 3)   # +3 per extra section, max +10
        skill_confidence[name] = min(100, round(best + freq_bonus))

    return skill_confidence

# ── Section detection (boolean flags) ──────────────────────────────────────────
def detect_sections(text: str) -> dict:
    lower = text.lower()
    found = {k: any(kw in lower for kw in kws) for k, kws in SECTION_KEYWORDS.items()}
    if not found['summary']:
        top = ' '.join(text.split()[:200])
        if any(len(s.split()) >= 15 for s in re.split(r'[.!?\n]', top)):
            found['summary'] = True
    return found

# ── Education level ─────────────────────────────────────────────────────────────
def detect_education_level(text: str) -> int:
    """Return numeric education level (1=school … 5=PhD)."""
    for pat, lvl in EDU_LEVELS:
        if pat.search(text):
            return lvl
    return 0

# ── Experience years ────────────────────────────────────────────────────────────
def extract_experience_years(text: str) -> float:
    """
    Parse experience years from multiple patterns.
    Returns total estimated years (capped at 30).
    """
    import datetime
    lower = text.lower()
    now = datetime.date.today()

    # Pattern 1: "X years of experience" / "X+ years"
    explicit = re.findall(r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s+(?:of\s+)?(?:experience|exp|work)', lower)
    if explicit:
        return min(30, max(float(v) for v in explicit))

    # Pattern 2: "experience of X years"
    exp2 = re.findall(r'experience\s+of\s+(\d+(?:\.\d+)?)\s*\+?\s*years?', lower)
    if exp2:
        return min(30, max(float(v) for v in exp2))

    # Pattern 3: month-year date ranges  e.g. Jan 2020 – Present
    months_map = {'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,
                  'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12}
    year_blocks = re.findall(
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})'
        r'\s*[-–—to/]+\s*'
        r'(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})|present|current|now|till\s+date)',
        lower
    )
    total_months = 0
    for blk in year_blocks:
        try:
            m1, y1 = months_map[blk[0]], int(blk[1])
            if blk[2]:
                m2, y2 = months_map[blk[2]], int(blk[3])
            else:
                m2, y2 = now.month, now.year
            total_months += max(0, (y2 - y1) * 12 + (m2 - m1))
        except Exception:
            pass
    if total_months:
        return min(30, round(total_months / 12, 1))

    # Pattern 4: year-only ranges  e.g. 2020 – 2023
    year_only = re.findall(
        r'\b(20\d{2})\s*[-–—to]+\s*(20\d{2}|present|current|now)\b', lower
    )
    total_months2 = 0
    for y1s, y2s in year_only:
        try:
            y1 = int(y1s)
            y2 = now.year if y2s in ('present','current','now') else int(y2s)
            total_months2 += max(0, (y2 - y1) * 12)
        except Exception:
            pass
    if total_months2:
        return min(30, round(total_months2 / 12, 1))

    # Pattern 5: bare numbers near year/yr
    nums = re.findall(r'(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:year|yr)s?', lower)
    valid = [float(n) for n in nums if 0 < float(n) <= 30]
    return max(valid) if valid else 0.0

# ── Certifications ──────────────────────────────────────────────────────────────
def detect_certifications(text: str) -> list:
    secs = split_sections(text)
    cert_text = secs.get('certifications', '') or text
    certs = []
    for line in cert_text.splitlines():
        line = line.strip()
        if line and CERT_RE.search(line) and len(line) > 5:
            certs.append(line[:80])
    return certs[:6]

# ── Identity extraction ─────────────────────────────────────────────────────────
_STOP = {'resume','cv','curriculum','vitae','profile','objective','summary',
         'linkedin','github','portfolio','contact','address','email','phone',
         'mobile','the','a','an','me','page','of','and','or','for'}

def extract_identity(text: str) -> dict:
    # Email
    em = re.search(r'[\w.+\-]+@[\w.\-]+\.[a-zA-Z]{2,}', text)
    email = em.group(0).lower() if em else '-'

    # Phone — try international then 10-digit
    phone = '-'
    for pat in (r'\+?(?:91|1|44)?[\s\-]?\(?\d{3,5}\)?[\s\-]?\d{3,5}[\s\-]?\d{3,5}',
                r'\b\d{10}\b'):
        m = re.search(pat, text)
        if m:
            d = re.sub(r'\D','', m.group(0))
            if len(d) >= 10:
                phone = d[-10:]
                break

    # LinkedIn
    linkedin = '-'
    lm = re.search(r'linkedin\.com/in/([\w\-]+)', text, re.I)
    if lm: linkedin = 'linkedin.com/in/' + lm.group(1)

    # GitHub
    github = '-'
    gm = re.search(r'github\.com/([\w\-]+)', text, re.I)
    if gm: github = 'github.com/' + gm.group(1)

    # Name — scan first 20 non-empty lines
    name = 'Unknown'
    checked = 0
    for line in text.splitlines():
        line = line.strip()
        if not line: continue
        checked += 1
        if checked > 20: break
        if re.search(r'[@\d\/\\\|*=<>{}()\[\]_+#]', line): continue
        if re.search(r'http|www\.|\.com|\.in\b', line, re.I): continue
        if len(line) > 55: continue
        words = line.split()
        if not (1 <= len(words) <= 5): continue
        if not all(w[0].isupper() for w in words if w and w[0].isalpha()): continue
        if any(w.lower() in _STOP for w in words): continue
        if sum(c.isalpha() for c in line) / max(len(line),1) < 0.72: continue
        name = line; break

    return {'name': name, 'email': email, 'phone': phone,
            'linkedin': linkedin, 'github': github}

# ── Scoring ─────────────────────────────────────────────────────────────────────
def compute_match_score(matched_skills: list, total_jd: int,
                        sections: dict, exp_years: float,
                        edu_level: int, certs: list,
                        skill_confidence: dict,
                        all_resume_skills: list = None) -> dict:
    """
    Multi-factor realistic scoring with strong differentiation:
      60 pts — Skill match (quadratic spread + confidence weighted)
      20 pts — Experience relevance
      10 pts — Education level
      10 pts — Profile completeness
    """
    # ── Skill component (60 pts) ────────────────────────────────────────────
    if total_jd <= 0:
        # No JD → score based on resume breadth
        breadth = len(all_resume_skills) if all_resume_skills else 0
        skill_comp = min(50.0, breadth * 2.5)
    else:
        n_matched = len(matched_skills)
        # Quadratic ratio: rewards more matches disproportionately
        base_ratio = (n_matched / total_jd) ** 0.75   # gentler curve than linear
        base_pts   = base_ratio * 52                    # up to 52 from ratio

        # Confidence boost (up to +8 pts)
        if matched_skills and skill_confidence:
            avg_conf = sum(skill_confidence.get(s, 40) for s in matched_skills) / len(matched_skills)
        else:
            avg_conf = 40
        conf_boost = (avg_conf / 100) * 8

        # Skill diversity bonus: more unique categories matched → +bonus
        matched_cats = set()
        for name, pat, cat, w in SKILL_PATTERNS:
            if name in matched_skills:
                matched_cats.add(cat)
        diversity_bonus = min(6, len(matched_cats) * 1.5)

        skill_comp = min(60.0, round(base_pts + conf_boost + diversity_bonus, 1))

    # ── Experience component (20 pts) — finer granularity ───────────────────
    if exp_years >= 10:   exp_comp = 20
    elif exp_years >= 7:  exp_comp = 18
    elif exp_years >= 5:  exp_comp = 16
    elif exp_years >= 3:  exp_comp = 13
    elif exp_years >= 2:  exp_comp = 10
    elif exp_years >= 1:  exp_comp = 8
    elif exp_years > 0:   exp_comp = 5
    elif sections.get('experience'): exp_comp = 4
    elif sections.get('projects'):   exp_comp = 3
    else:                            exp_comp = 1

    # ── Education component (10 pts) ────────────────────────────────────────
    edu_map = {5: 10, 4: 9, 3: 7, 2: 5, 1: 3, 0: 1}
    edu_comp = edu_map.get(edu_level, 1)

    # ── Profile completeness (10 pts) — penalise missing sections ───────────
    prof_comp = 0
    if sections.get('projects'):       prof_comp += 3
    if sections.get('skills'):         prof_comp += 2
    if sections.get('summary'):        prof_comp += 2
    if sections.get('certifications'): prof_comp += 2
    # Extra cert count bonus
    if len(certs) >= 3:                prof_comp = min(10, prof_comp + 1)

    total = round(skill_comp + exp_comp + edu_comp + prof_comp)
    total = max(8, min(97, total))

    return {
        'match_score':   total,
        'skill_comp':    round(skill_comp),
        'exp_comp':      exp_comp,
        'edu_comp':      edu_comp,
        'prof_comp':     prof_comp,
    }

def classify(score: int) -> str:
    if score >= 80: return 'Excellent'
    if score >= 65: return 'Good'
    if score >= 45: return 'Average'
    return 'Below Average'

def smart_cluster_label(resume_skills: set) -> str:
    """Label cluster based on dominant skill category."""
    cats = Counter()
    for name, pat, cat, w in SKILL_PATTERNS:
        if name in resume_skills:
            cats[cat] += w
    if not cats: return 'General Profile'
    top = cats.most_common(2)
    primary = top[0][0]
    secondary = top[1][0] if len(top) > 1 else None
    label_map = {
        ('ml', None):       'AI / ML Specialist',
        ('ml', 'pl'):       'ML Engineer',
        ('ml', 'be'):       'ML Data Engineer',
        ('fe', None):       'Frontend Developer',
        ('fe', 'be'):       'Full-Stack Developer',
        ('be', None):       'Backend Developer',
        ('be', 'cloud'):    'Cloud Backend Engineer',
        ('cloud', None):    'Cloud / DevOps Engineer',
        ('devops', None):   'DevOps Engineer',
        ('devops', 'cloud'):'Cloud DevOps Engineer',
        ('pl', None):       'Software Engineer',
        ('db', None):       'Database Engineer',
        ('design', None):   'UI/UX Designer',
        ('design', 'fe'):   'Frontend / UI Designer',
    }
    for (p, s), lbl in label_map.items():
        if primary == p and (s is None or secondary == s):
            return lbl
    return 'Software Engineer'

def build_recommendations(missing_skills, limit=8):
    recs = []
    for skill in missing_skills[:limit]:
        em, bg = SKILL_META.get(skill, (None,None))
        if not em:
            em, bg = _DM['emoji'], _DM['bg']
        q = skill.replace(' ','+').replace('/','')
        recs.append({
            'skill': skill, 'emoji': em, 'bg': bg,
            'youtube_url': f'https://www.youtube.com/results?search_query={q}+tutorial+for+beginners',
            'project_url': f'https://www.google.com/search?q={q}+project+ideas',
        })
    return recs

def validate_resume(text: str) -> bool:
    lower = text.lower()
    return sum(1 for kw in ['skill','education','experience','project'] if kw in lower) >= 2

# ── Fallback result ─────────────────────────────────────────────────────────────
def _fallback(filename, jd_skills, total_jd):
    return {
        'filename':filename,'name':'Unknown','email':'-','phone':'-',
        'linkedin':'-','github':'-',
        'match_score':10,'skill_score':0,'project_score':20,'exp_score':15,
        'candidate_exp':0,'has_edu':False,'edu_level':0,'edu_label':'Unknown',
        'certifications':[],'matched_skills':[],'missing_skills':sorted(jd_skills),
        'recommendations':build_recommendations(sorted(jd_skills)),
        'total_jd_skills':total_jd,'skills_list':[],'skill_confidence':{},
        'cluster':'General Profile','classification':'Below Average',
        'similar_candidates':[],'extraction_failed':True,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return send_from_directory(app.template_folder, 'index.html')

@app.route('/<path:filename>')
def static_files(filename):
    if filename.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(app.template_folder, filename)


@app.route('/api/analyze', methods=['POST'])
def analyze():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['resume']
    if not file.filename:
        return jsonify({'error': 'Empty filename'}), 400
    text = extract_text(file)
    if not text.strip():
        return jsonify({'error': 'Could not extract text. Ensure the PDF has selectable text.'}), 422
    if not validate_resume(text):
        return jsonify({'error': 'Upload a valid resume with Skills, Experience, Education, Projects.'}), 422

    identity   = extract_identity(text)
    sections   = detect_sections(text)
    skill_conf = detect_skills_sectioned(text)
    detected   = list(skill_conf.keys())
    exp_years  = extract_experience_years(text)
    edu_level  = detect_education_level(text)
    certs      = detect_certifications(text)

    important = ['python','javascript','typescript','react','node.js','sql','docker',
                 'git','aws','machine learning','rest api','system design','ci/cd',
                 'kubernetes','mongodb','postgresql','java','figma','graphql']
    missing    = [s for s in important if s not in skill_conf]
    scores     = compute_match_score(detected, len(important), sections,
                                     exp_years, edu_level, certs, skill_conf,
                                     all_resume_skills=detected)
    skill_score = min(100, round(len(detected) / len(important) * 100))
    proj_score  = (95 if sections.get('projects') and certs
                   else 80 if sections.get('projects') else 20)
    exp_score   = (min(98, int(exp_years / 12 * 100))
                   if exp_years >= 1 else (45 if sections.get('experience') else 12))
    edu_score   = edu_level * 18
    edu_labels  = {5:'PhD',4:'Masters',3:'Bachelors',2:'Diploma',1:'High School',0:'Unknown'}

    return jsonify({
        'resume_score':    scores['match_score'],
        'detected_skills': detected,
        'skill_confidence':skill_conf,
        'missing_skills':  missing,
        'missing_sections':[s for s,ok in sections.items() if not ok],
        'section_scores':  {
            'Skills & Technologies': skill_score,
            'Work Experience':       exp_score,
            'Education':             edu_score,
            'Projects':              proj_score,
            'Summary / Objective':   80 if sections.get('summary') else 20,
        },
        'recommendations':  build_recommendations(missing),
        'filename':         file.filename,
        'name':             identity['name'],
        'email':            identity['email'],
        'phone':            identity['phone'],
        'linkedin':         identity['linkedin'],
        'github':           identity['github'],
        'experience_years': exp_years,
        'edu_level':        edu_level,
        'edu_label':        edu_labels.get(edu_level, 'Unknown'),
        'certifications':   certs,
    })


@app.route('/api/match', methods=['POST'])
def match():
    print('[DEBUG] /api/match received')
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify({'error': 'No files uploaded'}), 400
        job_desc = request.form.get('job_description', '').strip()
        if not job_desc:
            return jsonify({'error': 'No job description provided'}), 400

        jd_skills = set(detect_skills(job_desc))
        total_jd  = len(jd_skills)
        print(f'[DEBUG] JD skills ({total_jd}): {sorted(jd_skills)}')

        results       = []
        valid_indices = []
        skill_vectors = []
        edu_labels    = {5:'PhD',4:'Masters',3:'Bachelors',2:'Diploma',1:'High School',0:'Unknown'}

        for file in files:
            if not file.filename:
                continue
            text = extract_text(file)
            print(f'[DEBUG] {file.filename}: {len(text)} chars')
            if not text.strip():
                results.append(_fallback(file.filename, jd_skills, total_jd))
                continue
            try:
                identity   = extract_identity(text)
                sections   = detect_sections(text)
                skill_conf = detect_skills_sectioned(text)
                resume_set = set(skill_conf.keys())
                matched    = sorted(resume_set & jd_skills)
                missing    = sorted(jd_skills - resume_set)
                exp_years  = extract_experience_years(text)
                edu_level  = detect_education_level(text)
                certs      = detect_certifications(text)
                sd         = compute_match_score(matched, total_jd, sections,
                                                 exp_years, edu_level, certs, skill_conf,
                                                 all_resume_skills=sorted(resume_set))
                match_score = sd['match_score']
                # Sub-scores shown on dashboard (normalised 0-100)
                skill_score = min(100, round(len(matched) / max(total_jd, 1) * 100)) if total_jd else min(100, len(resume_set) * 3)
                proj_score  = (95 if sections.get('projects') and len(certs) >= 2
                               else 80 if sections.get('projects')
                               else 35 if sections.get('experience') else 15)
                exp_score   = (min(98, int(exp_years / 12 * 100))
                               if exp_years >= 1 else (45 if sections.get('experience') else 12))
                has_edu     = bool(sections.get('education'))
                cls         = classify(match_score)
                cluster     = smart_cluster_label(resume_set)

                print(f'[DEBUG] {file.filename} | skills_found={len(resume_set)} | '
                      f'matched={len(matched)} | exp={exp_years}yr | '
                      f'edu_lvl={edu_level} | certs={len(certs)} | '
                      f'score={match_score} | cls={cls}')
                print(f'  -> matched_skills : {matched}')
                print(f'  -> top_skills     : {sorted(resume_set)[:12]}')
                print(f'  -> score_breakdown: skill={sd["skill_comp"]} '
                      f'exp={sd["exp_comp"]} edu={sd["edu_comp"]} '
                      f'prof={sd["prof_comp"]}')

                vector = [1 if n in resume_set else 0 for n,*_ in _SR]
                valid_indices.append(len(results))
                skill_vectors.append(vector)

                results.append({
                    'filename':        file.filename,
                    'name':            identity['name'],
                    'email':           identity['email'],
                    'phone':           identity['phone'],
                    'linkedin':        identity['linkedin'],
                    'github':          identity['github'],
                    'match_score':     match_score,
                    'skill_score':     skill_score,
                    'project_score':   proj_score,
                    'exp_score':       exp_score,
                    'candidate_exp':   exp_years,
                    'has_edu':         has_edu,
                    'edu_level':       edu_level,
                    'edu_label':       edu_labels.get(edu_level, 'Unknown'),
                    'certifications':  certs,
                    'matched_skills':  matched,
                    'missing_skills':  missing,
                    'recommendations': build_recommendations(missing),
                    'total_jd_skills': total_jd,
                    'skills_list':     sorted(resume_set),
                    'skill_confidence':skill_conf,
                    'cluster':         cluster,
                    'classification':  cls,
                    'similar_candidates': [],
                    'is_duplicate':    False,
                    'score_breakdown': sd,
                })
            except Exception as e:
                app.logger.warning(f'[MATCH] {file.filename}: {e}')
                results.append(_fallback(file.filename, jd_skills, total_jd))

        n_valid = len(valid_indices)
        print(f'[DEBUG] phase1 done: {len(results)} results, {n_valid} for ML')
        final_rules = []

        # KMeans
        if n_valid >= 2:
            try:
                va = np.array(skill_vectors, dtype=float)
                if not np.all(va.sum(axis=1) == 0):
                    k = min(4, n_valid)
                    KMeans(n_clusters=k, random_state=42, n_init='auto').fit_predict(va)
                app.logger.info('[KMEANS] done')
            except Exception as e:
                app.logger.warning(f'[KMEANS] {e}')

        # Apriori
        if apriori and n_valid >= 3:
            try:
                all_s = [results[ri]['skills_list'] for ri in valid_indices]
                cnts  = Counter(s for lst in all_s for s in lst)
                uniq  = sorted(s for s,_ in cnts.most_common(25))
                if len(uniq) >= 2:
                    df   = pd.DataFrame([{s:(s in c) for s in uniq} for c in all_s], dtype=bool)
                    freq = apriori(df, min_support=0.3, use_colnames=True)
                    if not freq.empty:
                        kw = dict(metric='confidence', min_threshold=0.7)
                        if _AR_NUM: kw['num_itemsets'] = len(df)
                        rules = association_rules(freq, **kw)
                        for _, row in rules.iterrows():
                            a  = list(row['antecedents'])[0]
                            c2 = list(row['consequents'])[0]
                            final_rules.append(
                                f"{a} \u2192 {c2} ({round(row['confidence']*100)}%)"
                            )
                app.logger.info(f'[APRIORI] {len(final_rules)} rules')
            except Exception as e:
                app.logger.warning(f'[APRIORI] {e}')

        # Cosine similarity + duplicate detection
        if n_valid >= 2:
            try:
                va  = np.array(skill_vectors, dtype=float)
                sim = cosine_similarity(va)
                for vi, ri in enumerate(valid_indices):
                    order   = sim[vi].argsort()[::-1]
                    similar = []
                    for si in order:
                        if si == vi: continue
                        if sim[vi][si] > 0.35:
                            similar.append(results[valid_indices[si]]['filename'])
                        if len(similar) >= 2: break
                    results[ri]['similar_candidates'] = similar
                    best = next((si for si in order if si != vi), None)
                    results[ri]['is_duplicate'] = bool(
                        best is not None and sim[vi][best] > 0.92
                    )
                app.logger.info('[COSINE] done')
            except Exception as e:
                app.logger.warning(f'[COSINE] {e}')

        if not results:
            for file in files:
                if file.filename:
                    results.append(_fallback(file.filename, jd_skills, total_jd))

        results.sort(key=lambda r: r['match_score'], reverse=True)
        print(f'[DEBUG] returning {len(results)} results')
        return jsonify({'results': results, 'job_description': job_desc,
                        'association_rules': final_rules[:5]})

    except Exception as e:
        print(f'[ERROR] {e}')
        app.logger.exception(f'[MATCH] {e}')
        return jsonify({'results': [], 'association_rules': [], 'error': str(e)}), 200


@app.errorhandler(404)
def not_found(e): return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(405)
def method_not_allowed(e): return jsonify({'error': 'Method not allowed'}), 405

@app.errorhandler(413)
def too_large(e): return jsonify({'error': 'File too large (max 16 MB)'}), 413

@app.errorhandler(500)
def server_error(e): return jsonify({'error': f'Internal server error: {e}'}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5000)